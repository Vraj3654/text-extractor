from fastapi import FastAPI, UploadFile, File, Depends, HTTPException, Form, Body, status
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from sqlalchemy.orm import Session
from database import SessionLocal, engine, Base
import models
import os
import shutil
from datetime import datetime, timedelta
import image_pipeline
import extract_text
import document_analysis
from pdf2image import convert_from_bytes
import numpy as np
import cv2
import json
import base64
from docx import Document as DocxDocument
from fpdf import FPDF

import bcrypt
from jose import JWTError, jwt
from pydantic import BaseModel

# Create database tables
Base.metadata.create_all(bind=engine)

app = FastAPI(title="Text Extractor API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

os.makedirs("static", exist_ok=True)
app.mount("/static", StaticFiles(directory="static"), name="static")

@app.on_event("startup")
def startup_event():
    try:
        extract_text.configure_tesseract()
        print("✅ Tesseract configured successfully.")
    except Exception as e:
        print(f"❌ Tesseract init error: {e}")

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# ===========================
# AUTHENTICATION
# ===========================
SECRET_KEY = "SUPER_SECRET_SAAS_KEY"
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24 * 7 # 1 week

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/api/token")

def verify_password(plain, hashed):
    return bcrypt.checkpw(plain.encode('utf-8'), hashed.encode('utf-8'))

def get_password_hash(password):
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def create_access_token(data: dict):
    to_encode = data.copy()
    expire = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

def get_current_user(token: str = Depends(oauth2_scheme), db: Session = Depends(get_db)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
    user = db.query(models.User).filter(models.User.username == username).first()
    if user is None:
        raise credentials_exception
    return user

class UserCreate(BaseModel):
    username: str
    password: str

@app.post("/api/register")
def register(user: UserCreate, db: Session = Depends(get_db)):
    db_user = db.query(models.User).filter(models.User.username == user.username).first()
    if db_user:
        raise HTTPException(status_code=400, detail="Username already registered")
    hashed_pw = get_password_hash(user.password)
    new_user = models.User(username=user.username, hashed_password=hashed_pw)
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    return {"message": "User created successfully"}

@app.post("/api/token")
def login(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    user = db.query(models.User).filter(models.User.username == form_data.username).first()
    if not user or not verify_password(form_data.password, user.hashed_password):
        raise HTTPException(status_code=400, detail="Incorrect username or password")
    access_token = create_access_token(data={"sub": user.username})
    return {"access_token": access_token, "token_type": "bearer"}
    
@app.get("/api/me")
def read_users_me(current_user: models.User = Depends(get_current_user)):
    return {
        "username": current_user.username
    }

# ===========================
# CORE ROUTES
# ===========================

@app.get("/")
def read_root():
    index_path = os.path.join("static", "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path)
    return {"message": "API is running. Place index.html in static/ directory to view UI."}

@app.post("/api/upload")
async def upload_image(
    file: UploadFile = File(...), 
    language: str = Form("eng"), 
    db: Session = Depends(get_db), 
    current_user: models.User = Depends(get_current_user)
):

    if not file.filename:
        raise HTTPException(status_code=400, detail="No selected file")
    
    try:
        file_bytes = await file.read()
        
        if file.filename.lower().endswith(".pdf"):
            pages = convert_from_bytes(file_bytes)
            if not pages:
                raise ValueError("PDF contains no pages")
            
            total_raw = ""
            total_corrected = ""
            confidences = []
            detected_langs = []
            
            for page in pages:
                open_cv_image = np.array(page) 
                open_cv_image = open_cv_image[:, :, ::-1].copy() # RGB to BGR
                processed_img_array = preprocess.preprocess_image(cv2.imencode('.jpg', open_cv_image)[1].tobytes())
                res = extract_text.extract_text_from_image(processed_img_array, languages=language)
                
                total_raw += res["raw_text"] + "\n\n"
                total_corrected += res["corrected_text"] + "\n\n"
                confidences.append(res["confidence"])
                if res["detected_language"] != "unknown":
                    detected_langs.append(res["detected_language"])
            
            final_raw = total_raw.strip()
            final_corrected = total_corrected.strip()
            final_confidence = sum(confidences) / len(confidences) if confidences else 0.0
            
            from collections import Counter
            final_lang = Counter(detected_langs).most_common(1)[0][0] if detected_langs else "unknown"
            
            result = {
                "raw_text": final_raw,
                "corrected_text": final_corrected,
                "confidence": final_confidence,
                "detected_language": final_lang
            }
        else:
            filename_lower = file.filename.lower()
            # Map shorthand language codes to Tesseract multi-language strings
            LANG_MAP = {
                "hin": "hin+eng",   # Aadhaar / Hindi docs need both scripts
                "guj": "guj+eng",
                "tel": "tel+eng",
                "tam": "tam+eng",
            }
            tess_lang = LANG_MAP.get(language, language)

            # General document upload always uses the scan pipeline
            processed_img_array = image_pipeline.process_scanned_document(file_bytes)
            result = extract_text.extract_text_from_image(processed_img_array, languages=tess_lang)
            
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
        
    new_doc = models.Document(
        filename=file.filename,
        raw_text=result.get("raw_text", ""),
        corrected_text=result.get("corrected_text", ""),
        language=result.get("detected_language", "unknown"),
        confidence=result.get("confidence", 0.0),
        user_id=current_user.id
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)
    
    return {
        "id": new_doc.id,
        "filename": new_doc.filename,
        "raw_text": new_doc.raw_text,
        "corrected_text": new_doc.corrected_text,
        "language": new_doc.language,
        "confidence": new_doc.confidence,
        "created_at": new_doc.created_at
    }

@app.post("/api/upload-id")
async def upload_id_document(
    file: UploadFile = File(...), 
    language: str = Form("eng"), 
    db: Session = Depends(get_db), 
    current_user: models.User = Depends(get_current_user)
):
    """Specific endpoint for Identity Documents (Smart Form feature)"""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No selected file")
    
    try:
        file_bytes = await file.read()
        
        # Map shorthand language codes to Tesseract multi-language strings
        LANG_MAP = {
            "hin": "hin+eng",
            "guj": "guj+eng",
            "tel": "tel+eng",
            "tam": "tam+eng",
        }
        tess_lang = LANG_MAP.get(language, language)

        # ID documents always use the gentler ID pipeline (preserves details of photos/cards)
        processed_img_array = image_pipeline.process_id_document(file_bytes)
        result = extract_text.extract_text_from_image(processed_img_array, languages=tess_lang)
        
        # Immediately run analysis for ID documents
        text = result.get("corrected_text", "") or result.get("raw_text", "")
        classification = document_analysis.classify_document(text)
        key_info = document_analysis.extract_key_info(text)

    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
        
    # Save the document history
    new_doc = models.Document(
        filename=file.filename,
        raw_text=result.get("raw_text", ""),
        corrected_text=result.get("corrected_text", ""),
        language=result.get("detected_language", "unknown"),
        confidence=result.get("confidence", 0.0),
        user_id=current_user.id
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)
    
    return {
        "id": new_doc.id,
        "filename": new_doc.filename,
        "raw_text": new_doc.raw_text,
        "corrected_text": new_doc.corrected_text,
        "language": new_doc.language,
        "confidence": new_doc.confidence,
        "classification": classification,
        "key_info": key_info
    }


@app.get("/api/history")
def get_history(db: Session = Depends(get_db), current_user: models.User = Depends(get_current_user)):
    docs = db.query(models.Document).filter(models.Document.user_id == current_user.id).order_by(models.Document.created_at.desc()).all()
    return docs

@app.get("/api/export/{doc_id}/{format_type}")
def export_document(doc_id: int, format_type: str, db: Session = Depends(get_db), current_user: models.User = Depends(get_current_user)):
    doc = db.query(models.Document).filter(models.Document.id == doc_id, models.Document.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
        
    text_to_export = doc.corrected_text if doc.corrected_text else doc.raw_text
    
    export_dir = "static/exports"
    os.makedirs(export_dir, exist_ok=True)
    
    base_name = os.path.splitext(doc.filename)[0]
    
    if format_type == "txt":
        file_path = os.path.join(export_dir, f"{base_name}.txt")
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(text_to_export)
        return FileResponse(path=file_path, filename=f"{base_name}.txt", media_type="text/plain")
        
    elif format_type == "json":
        file_path = os.path.join(export_dir, f"{base_name}.json")
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump({"filename": doc.filename, "text": text_to_export}, f, indent=4)
        return FileResponse(path=file_path, filename=f"{base_name}.json", media_type="application/json")
        
    elif format_type == "docx":
        file_path = os.path.join(export_dir, f"{base_name}.docx")
        d = DocxDocument()
        d.add_paragraph(text_to_export)
        d.save(file_path)
        return FileResponse(path=file_path, filename=f"{base_name}.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
    elif format_type == "pdf":
        file_path = os.path.join(export_dir, f"{base_name}.pdf")
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        safe_text = text_to_export.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 10, txt=safe_text)
        pdf.output(file_path)
        return FileResponse(path=file_path, filename=f"{base_name}.pdf", media_type="application/pdf")
    else:
        raise HTTPException(status_code=400, detail="Invalid format type")

@app.get("/api/analyze/{doc_id}")
def analyze_document(doc_id: int, db: Session = Depends(get_db), current_user: models.User = Depends(get_current_user)):
    doc = db.query(models.Document).filter(models.Document.id == doc_id, models.Document.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    text = doc.corrected_text or doc.raw_text or ""
    classification = document_analysis.classify_document(text)
    key_info = document_analysis.extract_key_info(text)

    return {
        "doc_id": doc_id,
        "filename": doc.filename,
        "classification": classification,
        "key_info": key_info
    }


@app.post("/api/translate")
async def translate_text(
    doc_id: int = Form(...),
    target_language: str = Form(...),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    doc = db.query(models.Document).filter(models.Document.id == doc_id, models.Document.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    text = doc.corrected_text or doc.raw_text or ""
    result = document_analysis.translate_text(text, target_language)

    if not result["success"]:
        raise HTTPException(status_code=500, detail=result.get("error", "Translation failed"))

    return result

@app.post("/api/camera-ocr")
async def camera_ocr(payload: dict = Body(...), db: Session = Depends(get_db), current_user: models.User = Depends(get_current_user)):

    try:
        data_url = payload.get("image", "")
        handwriting_mode = payload.get("handwriting", False)

        if not data_url:
            raise HTTPException(status_code=400, detail="No image data provided")

        if "," in data_url:
            data_url = data_url.split(",", 1)[1]

        img_bytes = base64.b64decode(data_url)

        import pytesseract

        # Use unified image pipeline for camera snapshots
        processed = image_pipeline.process_camera_snapshot(img_bytes, handwriting_mode)

        if handwriting_mode:
            # Try PSM 11 (sparse text) — best for handwriting
            psm11_text = pytesseract.image_to_string(processed, config='--oem 3 --psm 11').strip()
            # Try PSM 6 (single uniform block) as fallback
            psm6_text  = pytesseract.image_to_string(processed, config='--oem 3 --psm 6').strip()
            # Use whichever gives more output
            raw_text = psm11_text if len(psm11_text) >= len(psm6_text) else psm6_text
        else:
            raw_text = pytesseract.image_to_string(processed, config='--oem 3 --psm 3').strip()

        # Confidence score
        data = pytesseract.image_to_data(processed, config='--oem 3 --psm 3', output_type=pytesseract.Output.DICT)
        confidences = [int(c) for c in data['conf'] if str(c).lstrip('-').isdigit() and int(c) > 0]
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0

        return {
            "text": raw_text,
            "raw_text": raw_text,
            "confidence": round(avg_confidence, 1),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/supported-languages")
def get_supported_languages():
    return {"languages": document_analysis.SUPPORTED_LANGUAGES}


